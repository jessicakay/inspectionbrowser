import os
import csv
import re
import docx
import pandas

def gettext(report_files,sel):
    global logCSV, alltext
    global r, fileName
    global asteriskCount
    global totalPop, pop
    if sel=='a':
        print("loading...",prt)
        fileName = prt
    else:
        fileName = report_files[int(sel) - 1]
    fullPath = ''.join([filePath, fileName])
    doc = []
    repeatViolations=[]
    totalPop=[]
    report = docx.Document(fullPath)
    for paragraph in report.paragraphs:
        doc.append(paragraph.text)
        rp=re.search(r'[0-9]+\srepeat', str(paragraph.text))
        pp=re.search(r'population\sat\sthis\stime\swas\s[0-9]+', str(paragraph.text))
        p1=re.search(r'population\swas\s[0-9]+', str(paragraph.text))
        if rp != None:
            repeatViolations.append(str(paragraph.text))
        if pp !=None or p1 != None:
            totalPop.append(str(paragraph.text))
        elif pp==None and p1==None:
            totalPop.append("")


    strongs = re.findall(r'[0-9.-]+', str(doc))
    pop=re.findall(r'[0-9]+',str(totalPop))
    if pop == None:
        pop="Not found"
    pop=''.join(pop)
    print("pop #:",pop)

    asteriskCount = re.findall(r'\*', str(doc))
    r=re.findall(r'[0-9]+\srepeat\sviolations', str(repeatViolations))
    r=''.join(r)
    print(str(len(str(strongs))) + " strings evaluated")
    print("asterisks: "+str(len(asteriskCount)))
    print(str(r))

    logCSV = ''.join([filePath, "codearchive.csv"])
    alltext = []
    for strn in strongs:
        if "." in strn and re.search("[0-9]", strn):
            alltext.append(strn)
    getColHead(fullPath,alltext)

def getColHead(fullPath,alltext):
    doc = []
    report = docx.Document(fullPath)
    title  = report.core_properties.title
    print("the default header is ", title)
    print("the create data is: ", str(report.core_properties.created))
    thename=input("use this title? ")
    if thename == "y" or thename=="":
        title=re.sub(r"\s", "_", title)
        titleHead=[]
        for character in title:
            if character.isalnum()==True or character=='_':
                titleHead.append(str(character))
        titleHead=''.join(titleHead)
        head=re.sub("/","_",titleHead)
        head=str(head)
    else:
        newname=input("new column name: ")
        head=str(newname)
    print("using: ",head)
    writeSheet(head,fullPath,alltext)

def loader():
    global switch
    global i, data, fullPath, logCSV
    global desktopfiles, alltext, doc, strongs, home, filePath, report_files, sel
    switch=0
    home = (os.path.expanduser('~'))
    filePath = ''.join([home, "/Downloads/South Bay/"])
    os.chdir(filePath)
    desktopfiles = os.listdir(filePath)
    i = 1
    report_files = []
    for file in desktopfiles:
        if "~" not in file and "docx" in file:
            print(i, " - ", str(file))
            report_files.append(str(file))
            i += 1
    sel = input("number: ")
    if sel == 'u':
        print("coming soon!")
    if sel == 'a':
        allports(report_files)
    elif sel.isdigit()==True:
        if (int(sel)-1) < int(len(report_files)):
            gettext(report_files,sel)
        else:
            prompt=''.join(["number between 1 and ",str(len(report_files))])
            print(prompt)
            loader()
    else:
        print("invalid option")
        quit()

def allports(report_files):
    global prt
    fileNum = 0
    for prt in report_files:
        fullPath=prt
        sel="a"
        gettext(prt,sel)

def writeSheet(head,fullPath,alltext):
    global styledata, data
    global fullName
    global freq
    global report_name, create_date
    data=[]
    print("writesheet using: ", fileName)
    styledata = open(fullPath, 'rb')
    document = docx.Document(styledata)
    core_properties = document.core_properties
    create_date = str(core_properties.created)
    print("-> ",create_date)
    useDate=input("use this date?")
    if "n" in useDate:
        print("filename is: ",fileName)
        useDate=input("date override: ")
        create_date=str(useDate)
    report_name = str(core_properties.title)
    data=pandas.read_csv(logCSV,error_bad_lines=False)
    numCols=len(data.columns)
    newCol=pandas.Series([])
    i=1
    for i in range(len(alltext)):
        newCol[i]=alltext[i]
    data.insert(numCols,head,newCol,allow_duplicates=True)
    with open(logCSV,'r'):
        data.to_csv(logCSV,index=False,float_format='%g',header=True)
    print("Report created: " + create_date+" "+report_name+"\n")
    print(str(len(alltext))+" code violations found")
    print(str(numCols+1), " samples in dataset\ntop values: \n")
    styledata.close()
    freq=(newCol.value_counts([]))
    print(freq.head(5))
    freq=pandas.Series.to_frame(freq)
    buildDB(head)

def buildDB(head):
    freqT = pandas.DataFrame.transpose(freq)
    freqT.insert(1,"name",head)
    freqT.insert(1,"filename",str(fileName))
    print("pop = ",str(pop))
    if pop.isnumeric()==True:
        freqT.insert(1,"total_pop",int(pop))
    freqT.insert(1,"date",create_date)
    freqT.insert(1,"repeat_string",r)
    astc=int(str(len(asteriskCount)))
    freqT.insert(1,"ast_cnt",astc)
    print("\nfreq table extracted:\t\n", freqT)
    newSheet = ''.join([home,"/Downloads/South Bay/","sheet.csv"])
    ns_exists=os.path.exists(newSheet)
    if ns_exists==True:
        ns_file=pandas.read_csv('sheet.csv',sep=',')
        print("dataset loaded.")
        new=pandas.merge(ns_file,freqT,how='outer')
        new.to_csv('sheet.csv')
    elif ns_exists==False:
        print("no such file")
        freqT.to_csv('sheet.csv',index=False)

loader()